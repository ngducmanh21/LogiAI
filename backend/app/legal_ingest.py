"""Ingest văn bản luật upload lúc runtime -> cập nhật knowledge base.

Tái dùng logic chunk theo "Điều N." của scripts/clean_laws.py.
Chunks mới được append vào data/clean/legal_chunks.jsonl và (nếu có API key)
embed rồi append vào data/index/legal_embeddings.npy.
"""
import io
import json
import re
import unicodedata

from . import config

MAX_CHARS = 3500
MIN_CHARS = 40

ARTICLE_RE = re.compile(r"^\s*Điều\s+(\d+)\s*[\.:]?\s*(.*)", re.M)


def read_pdf_bytes(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def read_docx_bytes(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


def clean_text(t: str) -> str:
    t = unicodedata.normalize("NFC", t)
    t = t.replace("\u00a0", " ").replace("\t", " ")
    t = re.sub(r"[ ]{2,}", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def split_long(text: str) -> list[str]:
    if len(text) <= MAX_CHARS:
        return [text]
    parts = re.split(r"\n(?=\d{1,2}\.\s)", text)
    out, buf = [], ""
    for p in parts:
        if len(buf) + len(p) + 1 <= MAX_CHARS:
            buf = f"{buf}\n{p}".strip()
        else:
            if buf:
                out.append(buf)
            while len(p) > MAX_CHARS:
                cut = p.rfind(". ", 0, MAX_CHARS)
                cut = cut + 1 if cut > MAX_CHARS // 2 else MAX_CHARS
                out.append(p[:cut].strip())
                p = p[cut:].strip()
            buf = p
    if buf:
        out.append(buf)
    return out


def chunk_document(text: str, doc_id: str, title: str, source: str) -> list[dict]:
    chunks = []
    matches = list(ARTICLE_RE.finditer(text))
    if matches:
        pre = text[: matches[0].start()].strip()
        if len(pre) >= MIN_CHARS:
            for piece in split_long(pre):
                chunks.append({"article": None, "article_title": "Phần mở đầu/căn cứ",
                               "text": piece})
        for i, m in enumerate(matches):
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body = text[m.start():end].strip()
            if len(body) < MIN_CHARS:
                continue
            art_no = m.group(1)
            art_title = m.group(2).strip().rstrip(".")[:200]
            for piece in split_long(body):
                chunks.append({"article": art_no, "article_title": art_title,
                               "text": piece})
    else:
        for piece in split_long(text):
            chunks.append({"article": None, "article_title": None, "text": piece})

    out = []
    for k, c in enumerate(chunks):
        out.append({
            "chunk_id": f"{doc_id}#{k:03d}",
            "doc_id": doc_id,
            "doc_title": title,
            "article": c["article"],
            "article_title": c["article_title"],
            "text": c["text"],
            "source_file": source,
        })
    return out


def make_doc_id(name: str) -> str:
    stem = re.sub(r"\.(pdf|docx?)$", "", name, flags=re.I)
    return re.sub(r"[^A-Za-z0-9]+", "-", stem).strip("-")[:80] or "doc"


def ingest(data: bytes, filename: str, title: str | None = None) -> list[dict]:
    """Đọc file bytes -> list chunks. Raise ValueError nếu không đọc được."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        raw = read_pdf_bytes(data)
    elif lower.endswith(".docx"):
        raw = read_docx_bytes(data)
    else:
        raise ValueError("Chỉ hỗ trợ PDF hoặc DOCX.")
    text = clean_text(raw)
    if len(text) < 200:
        raise ValueError(
            f"Không trích xuất được văn bản (chỉ {len(text)} ký tự). "
            "File có thể là bản scan ảnh — cần bản có text layer.")
    doc_id = make_doc_id(filename)
    return chunk_document(text, doc_id, title or filename, f"uploads/{filename}")


def persist_chunks(chunks: list[dict]) -> None:
    """Append chunks vào data/clean/legal_chunks.jsonl."""
    path = config.DATA_CLEAN / "legal_chunks.jsonl"
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "a", encoding="utf-8") as f:
        for c in chunks:
            f.write(json.dumps(c, ensure_ascii=False) + "\n")

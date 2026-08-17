#!/usr/bin/env python3
"""Chunk văn bản luật trong data/Law -> data/clean/legal_chunks.jsonl.

- Ưu tiên .docx (python-docx); .doc/.pdf dùng bản .pdf (pypdf).
- Chunk theo cấu trúc "Điều N." ; điều quá dài cắt tiếp theo khoản/độ dài.
Chạy:  python scripts/clean_laws.py
"""
import json
import re
import sys
import unicodedata
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
LAW_DIR = ROOT / "data" / "Law"
OUT = ROOT / "data" / "clean" / "legal_chunks.jsonl"

MAX_CHARS = 3500          # ~ giới hạn 1 chunk
MIN_CHARS = 40


def read_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def pick_sources(folder: Path) -> list[Path]:
    """Chọn 1 file tốt nhất cho mỗi văn bản trong folder (docx > pdf)."""
    files = sorted(folder.iterdir())
    chosen: dict[str, Path] = {}
    for f in files:
        if f.suffix.lower() == ".docx":
            chosen[f.stem] = f
    for f in files:
        if f.suffix.lower() == ".pdf":
            # "x.docx.pdf" trùng với "x.docx" -> bỏ nếu đã có docx
            base = f.stem  # "x.docx" hoặc "x"
            base = re.sub(r"\.docx?$", "", base)
            if not any(re.sub(r"\.docx?$", "", k) == base for k in chosen):
                chosen[f.stem] = f
    return list(chosen.values())


DOC_NO_RE = re.compile(r"(\d+[_/-]\d{4}[_/-][A-ZĐ]+(?:[_-][A-ZĐ]+)*|\d{4,6}[_ -](?:TB|CV|CHQ)[\w-]*)", re.I)
ARTICLE_RE = re.compile(r"^\s*Điều\s+(\d+)\s*[\.:]?\s*(.*)", re.M)
CHAPTER_RE = re.compile(r"^\s*(Chương\s+[IVXLC\d]+.*)$", re.M)


def clean_text(t: str) -> str:
    t = unicodedata.normalize("NFC", t)
    t = t.replace("\u00a0", " ").replace("\t", " ")
    t = re.sub(r"[ ]{2,}", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def split_long(text: str) -> list[str]:
    """Cắt đoạn dài theo khoản (1., 2., a) ...) hoặc theo câu."""
    if len(text) <= MAX_CHARS:
        return [text]
    # thử cắt theo khoản "n."
    parts = re.split(r"\n(?=\d{1,2}\.\s)", text)
    out, buf = [], ""
    for p in parts:
        if len(buf) + len(p) + 1 <= MAX_CHARS:
            buf = f"{buf}\n{p}".strip()
        else:
            if buf:
                out.append(buf)
            while len(p) > MAX_CHARS:          # đoạn đơn quá dài -> cắt cứng theo câu
                cut = p.rfind(". ", 0, MAX_CHARS)
                cut = cut + 1 if cut > MAX_CHARS // 2 else MAX_CHARS
                out.append(p[:cut].strip())
                p = p[cut:].strip()
            buf = p
    if buf:
        out.append(buf)
    return out


def chunk_document(text: str, doc_id: str, title: str, source: str) -> list[dict]:
    chunks = []
    matches = list(ARTICLE_RE.finditer(text))
    if matches:
        # preamble (tên văn bản, căn cứ) trước Điều 1
        pre = text[: matches[0].start()].strip()
        if len(pre) >= MIN_CHARS:
            for j, piece in enumerate(split_long(pre)):
                chunks.append({"article": None, "article_title": "Phần mở đầu/căn cứ",
                               "text": piece, "part": j})
        for i, m in enumerate(matches):
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            body = text[m.start():end].strip()
            if len(body) < MIN_CHARS:
                continue
            art_no = m.group(1)
            art_title = m.group(2).strip().rstrip(".")[:200]
            for j, piece in enumerate(split_long(body)):
                chunks.append({"article": art_no, "article_title": art_title,
                               "text": piece, "part": j})
    else:
        # không có cấu trúc Điều (công văn, thông báo) -> cắt theo độ dài
        for j, piece in enumerate(split_long(text)):
            chunks.append({"article": None, "article_title": None, "text": piece, "part": j})

    out = []
    for k, c in enumerate(chunks):
        out.append({
            "chunk_id": f"{doc_id}#{k:03d}",
            "doc_id": doc_id,
            "doc_title": title,
            "article": c["article"],
            "article_title": c["article_title"],
            "text": c["text"],
            "source_file": source,
        })
    return out


def main():
    all_chunks = []
    docs = []
    for folder in sorted(LAW_DIR.iterdir()):
        if not folder.is_dir():
            continue
        title = folder.name.strip()
        for src in pick_sources(folder):
            try:
                raw = read_docx(src) if src.suffix.lower() == ".docx" else read_pdf(src)
            except Exception as e:
                print(f"LỖI đọc {src.name}: {e}", file=sys.stderr)
                continue
            text = clean_text(raw)
            if len(text) < 200:
                print(f"BỎ QUA (quá ngắn/scan ảnh): {src.name} ({len(text)} ký tự)",
                      file=sys.stderr)
                continue
            m = DOC_NO_RE.search(src.stem)
            doc_no = (m.group(1).replace("_", "/") if m else src.stem)
            doc_id = re.sub(r"[^A-Za-z0-9]+", "-", doc_no).strip("-")
            cks = chunk_document(text, doc_id, title, str(src.relative_to(ROOT)))
            all_chunks.extend(cks)
            docs.append((doc_id, title, src.name, len(text), len(cks)))
            print(f"OK  {doc_id:<24} {len(cks):>3} chunks  ({src.name})")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with open(OUT, "w", encoding="utf-8") as f:
        for c in all_chunks:
            f.write(json.dumps(c, ensure_ascii=False) + "\n")
    print(f"\nTổng: {len(docs)} văn bản, {len(all_chunks)} chunks -> {OUT}")


if __name__ == "__main__":
    main()
